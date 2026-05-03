"""
Servicio de generacion de oficios en formato DOCX.

Genera documentos oficiales de la DPP con formato institucional
del Gobierno del Estado de Michoacan, basado en el modelo de oficio
oficial proporcionado.
"""
import io
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree as _lxml_etree

STATIC_DIR = Path(__file__).parent.parent / "static"
LOGOS_DIR = STATIC_DIR / "logos"
FIRMAS_DIR = STATIC_DIR / "firmas"

GUINDA = RGBColor(0x91, 0x1A, 0x3A)

MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


class OficioGeneratorService:
    """Genera documentos DOCX institucionales para la DPP."""

    def generar_oficio_respuesta(
        self,
        *,
        folio_respuesta: str,
        fecha_respuesta: str,
        lugar: str = "Morelia, Michoacán",
        destinatario_nombre: str,
        destinatario_cargo: str,
        destinatario_dependencia: str,
        seccion_fundamento: str,
        seccion_referencia: str,
        seccion_objeto: str,
        seccion_cierre: str,
        firmante_nombre: str = "Mtro. Marco Antonio Flores Mejía",
        firmante_cargo: str = "Director de Programación y Presupuesto",
        referencia_elaboro: Optional[str] = None,
        referencia_reviso: Optional[str] = None,
        copias: Optional[list[str]] = None,
        incluir_firma_visual: bool = False,
        sello_digital_data: Optional[dict] = None,
        asunto: Optional[str] = None,
        tabla_imagen_path: Optional[str] = None,
        tabla_datos_json: Optional[list] = None,
    ) -> bytes:
        """
        Genera un oficio de respuesta completo en formato DOCX.
        Incluye recuadro institucional superior derecho.
        Retorna los bytes del archivo .docx.
        """
        doc = Document()

        self._set_page_margins(doc)
        self._set_default_font(doc)

        # ── Detectar membrete activo ────────────────────────────────────────
        from app.services.oficio_pdf_service import _get_membrete_activo, _get_membrete_config
        membrete_path = _get_membrete_activo()
        membrete_activo = membrete_path is not None

        if membrete_activo:
            # ── Con membrete: fondo PNG de página completa ──────────────────
            self._add_membrete_background_docx(doc, membrete_path)
            # ── Espaciador preciso para que la fecha quede justo debajo ─────
            # Página carta = 792 pt; fecha_y (desde abajo) → desde arriba = 792 - fecha_y
            # Margen superior DOCX = 1.5 cm = 42.52 pt
            # Espacio en área de contenido = (792 - fecha_y) - 42.52
            cfg = _get_membrete_config()
            fecha_y     = cfg.get("fecha_y", 620)
            _TOP_MARGIN = 1.5 * 28.3465          # ≈ 42.52 pt
            space_pt    = max((792 - fecha_y) - _TOP_MARGIN, 1.0)
            p_space = doc.add_paragraph()
            p_space.paragraph_format.space_before = Pt(0)
            p_space.paragraph_format.space_after  = Pt(0)
            p_space.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_space.paragraph_format.line_spacing  = Pt(space_pt)
            p_space.add_run().font.size = Pt(1)   # run mínimo para anclar la altura
        else:
            # ── Sin membrete: encabezado institucional estándar ─────────────
            self._add_identidad_institucional(doc)
            self._add_recuadro_institucional(
                doc,
                folio=folio_respuesta,
                asunto_corto=self._truncar_asunto(asunto or "El que se indica"),
            )
            self._add_lema(doc)

        # Fecha (sin space_before extra cuando membrete ya gestionó el espaciado)
        self._add_fecha(doc, lugar, fecha_respuesta,
                        space_before_pt=0 if membrete_activo else 8)
        self._add_empty_lines(doc, 1)
        self._add_destinatario(doc, destinatario_nombre, destinatario_cargo, destinatario_dependencia)
        self._add_empty_lines(doc, 1)

        # Cuerpo: 4 secciones
        body_text = "\n\n".join(
            s for s in [seccion_fundamento, seccion_referencia, seccion_objeto, seccion_cierre]
            if s and s.strip()
        )
        if body_text:
            self._add_body_text(doc, body_text, tabla_imagen_path=tabla_imagen_path, tabla_datos_json=tabla_datos_json)

        self._add_empty_lines(doc, 1)
        self._add_atentamente(doc)

        # Espacio para firma visual o espacio en blanco
        if sello_digital_data is None:
            if incluir_firma_visual:
                firma_path = FIRMAS_DIR / "firma_mafm.png"
                if firma_path.exists():
                    pf = doc.add_paragraph()
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pf.paragraph_format.space_after = Pt(0)
                    run_img = pf.add_run()
                    run_img.add_picture(str(firma_path), width=Inches(2))
                else:
                    self._add_empty_lines(doc, 4)
            else:
                self._add_empty_lines(doc, 4)
        else:
            self._add_empty_lines(doc, 2)

        # Nombre + Cargo SIEMPRE antes del sello
        self._add_firmante_nombre(doc, firmante_nombre, firmante_cargo)

        # Sello digital DESPUÉS del nombre (si aplica)
        if sello_digital_data is not None:
            self._add_sello_digital(
                doc,
                qr_png_bytes=sello_digital_data.get("qr_png_bytes", b""),
                nombre_firmante=sello_digital_data.get("nombre_firmante", ""),
                cargo_firmante=sello_digital_data.get("cargo_firmante", "Director de Programación y Presupuesto"),
                rfc_firmante=sello_digital_data.get("rfc_firmante", ""),
                serial_certificado=sello_digital_data.get("serial_certificado", ""),
                fecha_firma=sello_digital_data.get("fecha_firma", ""),
                correo_firmante=sello_digital_data.get("correo_firmante", ""),
                folio_firma=sello_digital_data.get("folio_firma", ""),
            )

        if copias:
            self._add_empty_lines(doc, 2)
            self._add_copias(doc, copias)

        if referencia_elaboro or referencia_reviso:
            self._add_referencia(doc, referencia_elaboro, referencia_reviso)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ------- Internos -------

    def _add_membrete_background_docx(self, doc: Document, image_path: str) -> None:
        """Inserta el membrete PNG como imagen de fondo de página completa (detrás del texto)."""
        from docx.opc.part import Part as _OpcPart
        from docx.opc.packuri import PackURI as _PackURI

        # 1. Leer imagen y registrar como relación del documento
        ext = Path(image_path).suffix.lower()
        ct  = "image/png" if ext == ".png" else "image/jpeg"
        with open(image_path, "rb") as f:
            img_bytes = f.read()

        image_part = _OpcPart(
            partname=_PackURI(f"/word/media/membrete_bg{ext}"),
            content_type=ct,
            blob=img_bytes,
        )
        rId = doc.part.relate_to(
            image_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        )

        # 2. Dimensiones en EMU (página carta: 8.5" × 11" = 7772400 × 10058400 EMU)
        cx, cy = 7772400, 10058400

        # 3. XML del anchor (imagen detrás del texto, posición absoluta 0,0 desde la página)
        NS = {
            "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
            "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        ns_decl = " ".join(f'xmlns:{k}="{v}"' for k, v in NS.items())

        xml_str = f"""<w:p {ns_decl}>
  <w:r>
    <w:drawing>
      <wp:anchor distT="0" distB="0" distL="0" distR="0"
                 simplePos="0" relativeHeight="251658240" behindDoc="1"
                 locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="900" name="MembreteFondo"/>
        <wp:cNvGraphicFramePr>
          <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="{NS['pic']}">
            <pic:pic>
              <pic:nvPicPr>
                <pic:cNvPr id="900" name="MembreteFondo"/>
                <pic:cNvPicPr/>
              </pic:nvPicPr>
              <pic:blipFill>
                <a:blip r:embed="{rId}"/>
                <a:stretch><a:fillRect/></a:stretch>
              </pic:blipFill>
              <pic:spPr>
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              </pic:spPr>
            </pic:pic>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>"""

        # 4. Insertar al inicio del body
        p_elem = _lxml_etree.fromstring(xml_str.encode("utf-8"))
        doc.element.body.insert(0, p_elem)

    def _set_page_margins(self, doc: Document) -> None:
        """Pagina carta con margenes del modelo oficial."""
        for section in doc.sections:
            section.page_width = Emu(7772400)     # ~21.6cm (carta)
            section.page_height = Emu(10058400)   # ~26.4cm (carta)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.2)

    def _set_default_font(self, doc: Document) -> None:
        """Establece la fuente por defecto."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

    def _add_identidad_institucional(self, doc: Document) -> None:
        """
        Agrega la identidad institucional del Gobierno del Estado.
        Formato oficial: Escudo nacional + "Gobierno del Estado de Michoacán de Ocampo"
        alineado a la izquierda, sin líneas decorativas ni colores.
        """
        # Escudo nacional (si existe)
        escudo_path = LOGOS_DIR / "escudo_mich.png"
        if escudo_path.exists():
            pe = doc.add_paragraph()
            pe.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pe.paragraph_format.space_after = Pt(2)
            pe.paragraph_format.space_before = Pt(0)
            run_e = pe.add_run()
            run_e.add_picture(str(escudo_path), height=Cm(1.8))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run("Gobierno del Estado\nde Michoacán de Ocampo")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _add_header_image(self, doc: Document) -> None:
        """Agrega la imagen del membrete institucional (incluye recuadro vacío).
        NOTA: No usar junto con _add_recuadro_institucional() para evitar duplicar recuadro.
        """
        header_path = LOGOS_DIR / "header_dpp.png"
        if not header_path.exists():
            self._add_header_text(doc)
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        # Ancho de contenido: ~16.6cm (pagina 21.6 - margenes 5.0)
        run.add_picture(str(header_path), width=Cm(16.6))

    def _add_separator_line(self, doc: Document) -> None:
        """Agrega linea separadora guinda."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        # Linea mediante border bottom en el parrafo
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "12",
                qn("w:space"): "1",
                qn("w:color"): "911A3A",
            },
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_recuadro_institucional(
        self, doc: Document,
        folio: str,
        asunto_corto: str = "El que se indica",
        oficina: str = "Dirección de Programación y Presupuesto",
    ) -> None:
        """
        Agrega el recuadro institucional superior derecho.
        6 filas × 2 columnas: etiqueta | valor.
        Formato del modelo oficial de la DPP.
        """
        rows_data = [
            ("Dependencia:", "Secretaría de Finanzas y Administración"),
            ("Sub-dependencia:", "Subsecretaría de Finanzas"),
            ("Oficina:", oficina),
            ("No. de oficio:", folio or "—"),
            ("Expediente:", "General"),
            ("Asunto:", asunto_corto),
        ]

        table = doc.add_table(rows=len(rows_data), cols=2)
        table.autofit = False

        # Alinear tabla a la derecha
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
        jc = tbl_pr.makeelement(qn("w:jc"), {qn("w:val"): "right"})
        tbl_pr.append(jc)

        # Ancho total ~8cm: etiqueta 2.5cm + valor 5.5cm
        table.columns[0].width = Cm(2.5)
        table.columns[1].width = Cm(5.5)

        # Bordes sutiles
        tbl_borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_el = tbl_borders.makeelement(
                qn(f"w:{border_name}"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "0",
                    qn("w:color"): "999999",
                },
            )
            tbl_borders.append(border_el)
        tbl_pr.append(tbl_borders)
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)

        # Llenar datos
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]

            # Celda etiqueta
            cell_label = row.cells[0]
            cell_label.width = Cm(2.5)
            p_label = cell_label.paragraphs[0]
            p_label.paragraph_format.space_after = Pt(0)
            p_label.paragraph_format.space_before = Pt(0)
            run_label = p_label.add_run(label)
            run_label.bold = True
            run_label.font.name = "Arial"
            run_label.font.size = Pt(7)
            run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Celda valor
            cell_value = row.cells[1]
            cell_value.width = Cm(5.5)
            p_value = cell_value.paragraphs[0]
            p_value.paragraph_format.space_after = Pt(0)
            p_value.paragraph_format.space_before = Pt(0)
            run_value = p_value.add_run(value)
            run_value.font.name = "Arial"
            run_value.font.size = Pt(7)
            run_value.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Padding mínimo en celdas
            for cell in (cell_label, cell_value):
                tc_pr = cell._tc.get_or_add_tcPr()
                mar = tc_pr.makeelement(qn("w:tcMar"), {})
                for side in ("top", "bottom", "start", "end"):
                    side_el = mar.makeelement(
                        qn(f"w:{side}"),
                        {qn("w:w"): "30", qn("w:type"): "dxa"},
                    )
                    mar.append(side_el)
                tc_pr.append(mar)

    def _add_lema(self, doc: Document) -> None:
        """Agrega el lema institucional del año, centrado entre comillas."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('"50 Aniversario de los Santuarios de la Mariposa Monarca"')
        run.font.size = Pt(10)
        run.italic = True

    def _add_fecha(self, doc: Document, lugar: str, fecha: str,
                   space_before_pt: float = 8) -> None:
        """Solo la línea de fecha, alineada a la derecha."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(space_before_pt)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{lugar}, a {fecha}.")
        run.font.size = Pt(11)

    @staticmethod
    def _truncar_asunto(asunto: str, max_chars: int = 60) -> str:
        """Trunca el asunto para que quepa en el recuadro."""
        if len(asunto) <= max_chars:
            return asunto
        return asunto[:max_chars - 3].rstrip() + "..."

    def _add_folio_fecha(self, doc: Document, folio: str, lugar: str, fecha: str) -> None:
        """Bloque de folio y fecha, alineado a la derecha (legacy, mantenido por compatibilidad)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"OFICIO No. {folio}")
        run.bold = True
        run.font.size = Pt(11)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run(f"{lugar}, {fecha}")
        run2.font.size = Pt(11)

    def _add_empty_lines(self, doc: Document, count: int) -> None:
        """Agrega lineas vacias."""
        for _ in range(count):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    def _add_destinatario(
        self, doc: Document,
        nombre: str, cargo: str, dependencia: str,
    ) -> None:
        """Bloque de destinatario. Solo incluye líneas con contenido real."""
        lines = []
        if nombre and nombre.strip() not in ("", "---"):
            lines.append(nombre.upper())
        if cargo and cargo.strip() not in ("", "---"):
            lines.append(cargo.upper())
        if dependencia and dependencia.strip() not in ("", "---"):
            dep = dependencia.upper()
            if not dep.endswith('.'):
                dep += '.'
            lines.append(dep)
        lines.append("PRESENTE.")
        for text in lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)

    @staticmethod
    def _is_table_line(line: str) -> bool:
        """Detecta si una línea es parte de una tabla markdown (contiene |)."""
        stripped = line.strip()
        return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2

    @staticmethod
    def _is_separator_line(line: str) -> bool:
        """Detecta líneas separadoras de tabla como |---|---|"""
        stripped = line.strip().replace(" ", "")
        if not stripped.startswith("|"):
            return False
        # Líneas tipo |---|---| o |:---|:---| etc
        inner = stripped.strip("|")
        parts = inner.split("|")
        return all(set(p.strip()).issubset({"-", ":", " "}) and len(p.strip()) > 0 for p in parts)

    def _parse_table_cells(self, line: str) -> list[str]:
        """Extrae las celdas de una línea de tabla markdown."""
        stripped = line.strip()
        if stripped.startswith("|"):
            stripped = stripped[1:]
        if stripped.endswith("|"):
            stripped = stripped[:-1]
        return [cell.strip() for cell in stripped.split("|")]

    def _add_word_table(self, doc: Document, table_lines: list[str]) -> None:
        """Convierte líneas de tabla markdown en una tabla Word formateada."""
        import re
        # Filtrar líneas separadoras (|---|---|)
        data_lines = [l for l in table_lines if not self._is_separator_line(l)]
        if not data_lines:
            return

        # Parsear todas las filas
        rows_data = [self._parse_table_cells(line) for line in data_lines]
        if not rows_data:
            return

        # Determinar número de columnas (máximo entre todas las filas)
        num_cols = max(len(row) for row in rows_data)
        num_rows = len(rows_data)

        # Crear tabla Word
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.autofit = True

        # Estilo de bordes
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_el = tblBorders.makeelement(
                qn(f"w:{border_name}"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "0",
                    qn("w:color"): "333333",
                },
            )
            tblBorders.append(border_el)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            is_header = (i == 0)
            for j in range(num_cols):
                cell = row.cells[j]
                cell_text = row_data[j] if j < len(row_data) else ""
                # Limpiar marcadores markdown de bold
                cell_text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)

                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = p.add_run(cell_text_clean)
                run.font.name = "Arial"
                run.font.size = Pt(9)

                if is_header or cell_text != cell_text_clean:
                    run.bold = True
                    # Fondo gris claro para encabezado
                    if is_header:
                        shading = cell._tc.get_or_add_tcPr().makeelement(
                            qn("w:shd"),
                            {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "E8E8E8"},
                        )
                        cell._tc.get_or_add_tcPr().append(shading)

                # Padding de celdas
                tc_pr = cell._tc.get_or_add_tcPr()
                mar = tc_pr.makeelement(qn("w:tcMar"), {})
                for side in ("top", "bottom", "start", "end"):
                    side_el = mar.makeelement(
                        qn(f"w:{side}"),
                        {qn("w:w"): "40", qn("w:type"): "dxa"},
                    )
                    mar.append(side_el)
                tc_pr.append(mar)

    def _preprocess_tabla_blocks(self, text: str) -> str:
        """Convierte bloques [TABLA]...[/TABLA] a formato markdown con pipes."""
        import re
        def _convert_block(match: re.Match) -> str:
            block = match.group(1).strip()
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                return ""
            # Detectar si las líneas ya tienen formato pipe
            has_pipes = any("|" in l for l in lines)
            if has_pipes:
                # Convertir líneas tipo "1 | dato" a "| 1 | dato |"
                result = []
                for line in lines:
                    parts = [p.strip() for p in line.split("|")]
                    # Asegurar que empiece y termine con |
                    result.append("| " + " | ".join(parts) + " |")
                # Agregar separador después del header
                if len(result) > 1:
                    num_cols = result[0].count("|") - 1
                    sep = "|" + "|".join(["---"] * num_cols) + "|"
                    result.insert(1, sep)
                return "\n".join(result)
            return block
        return re.sub(r'\[TABLA\](.*?)\[/TABLA\]', _convert_block, text, flags=re.DOTALL)

    @staticmethod
    def _has_pipe_content(line: str) -> bool:
        """Detecta si una línea contiene datos tabulares con | (no necesita empezar con |)."""
        stripped = line.strip()
        # Debe tener al menos un | y no ser solo separadores
        if "|" not in stripped:
            return False
        parts = stripped.split("|")
        # Al menos 2 partes con contenido
        non_empty = [p.strip() for p in parts if p.strip()]
        return len(non_empty) >= 2

    def _add_body_text(self, doc: Document, text: str, *, tabla_imagen_path: Optional[str] = None, tabla_datos_json: Optional[list] = None) -> None:
        """Agrega el cuerpo del oficio con formato justificado.
        Detecta tablas markdown y bloques [TABLA]...[/TABLA] y los convierte en tablas Word.
        Si hay tabla_datos_json, inserta una tabla Word real con esos datos.
        Si hay tabla_imagen_path, inserta la imagen en lugar de la primera tabla detectada.
        IMPORTANTE: Si se subió tabla (Excel o imagen) y el texto de la IA no contiene
        marcadores de tabla, la tabla se inserta al final del cuerpo automáticamente.
        """
        # Preprocesar: convertir [TABLA]...[/TABLA] a formato markdown pipe
        text = self._preprocess_tabla_blocks(text)

        tabla_imagen_usada = False
        tabla_excel_usada = False
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Detectar inicio de tabla (líneas con | como tabla)
            if self._is_table_line(line) or (self._has_pipe_content(line) and not line.strip().startswith("[")):
                table_lines = []
                while i < len(lines) and (
                    self._is_table_line(lines[i])
                    or self._is_separator_line(lines[i])
                    or self._has_pipe_content(lines[i])
                ):
                    table_lines.append(lines[i])
                    i += 1

                # Prioridad: 1) datos Excel, 2) imagen, 3) tabla markdown generada por IA
                if tabla_datos_json and not tabla_excel_usada:
                    self._add_excel_table(doc, tabla_datos_json)
                    tabla_excel_usada = True
                elif tabla_imagen_path and not tabla_imagen_usada:
                    self._add_table_image(doc, tabla_imagen_path)
                    tabla_imagen_usada = True
                else:
                    # Normalizar: asegurar que todas empiecen/terminen con |
                    normalized = []
                    for tl in table_lines:
                        s = tl.strip()
                        if not s.startswith("|"):
                            s = "| " + s
                        if not s.endswith("|"):
                            s = s + " |"
                        normalized.append(s)
                    self._add_word_table(doc, normalized)
                continue

            # Saltar líneas separadoras sueltas
            if self._is_separator_line(line):
                i += 1
                continue

            # Línea vacía = salto de párrafo
            stripped = line.strip()
            if not stripped:
                i += 1
                continue

            # Acumular párrafo (líneas consecutivas no-tabla, no-vacías)
            para_lines = []
            while i < len(lines):
                l = lines[i].strip()
                if not l or self._is_table_line(lines[i]) or self._is_separator_line(lines[i]) or self._has_pipe_content(lines[i]):
                    break
                para_lines.append(l)
                i += 1

            paragraph_text = " ".join(para_lines)
            if paragraph_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.first_line_indent = Cm(1.0)
                run = p.add_run(paragraph_text)
                run.font.name = "Arial"
                run.font.size = Pt(11)

        # ── GARANTÍA: Si se subió tabla y NO se insertó durante el recorrido,
        #    insertarla al final del cuerpo (después del último párrafo) ──
        if tabla_datos_json and not tabla_excel_usada:
            self._add_excel_table(doc, tabla_datos_json)
        elif tabla_imagen_path and not tabla_imagen_usada:
            self._add_table_image(doc, tabla_imagen_path)

    def _add_table_image(self, doc: Document, image_path: str) -> None:
        """Inserta una imagen de tabla/cuadro en el documento."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        # Ancho máximo del contenido (página menos márgenes)
        run.add_picture(image_path, width=Cm(14.5))

    def _add_excel_table(self, doc: Document, table_data: list[list[str]]) -> None:
        """Inserta una tabla Word real a partir de datos extraídos de Excel.
        table_data: lista de filas, cada fila es una lista de strings.
        La primera fila se trata como encabezado.
        """
        if not table_data or not table_data[0]:
            return

        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.autofit = True

        # Bordes
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_el = tblBorders.makeelement(
                qn(f"w:{border_name}"),
                {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): "333333"},
            )
            tblBorders.append(border_el)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            is_header = (i == 0)
            for j in range(num_cols):
                cell = row.cells[j]
                cell_text = row_data[j] if j < len(row_data) else ""

                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = p.add_run(cell_text)
                run.font.name = "Arial"
                run.font.size = Pt(9)

                if is_header:
                    run.bold = True
                    shading = cell._tc.get_or_add_tcPr().makeelement(
                        qn("w:shd"),
                        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "E8E8E8"},
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

                # Padding
                tc_pr = cell._tc.get_or_add_tcPr()
                mar = tc_pr.makeelement(qn("w:tcMar"), {})
                for side in ("top", "bottom", "start", "end"):
                    side_el = mar.makeelement(
                        qn(f"w:{side}"),
                        {qn("w:w"): "40", qn("w:type"): "dxa"},
                    )
                    mar.append(side_el)
                tc_pr.append(mar)

    def _add_atentamente(self, doc: Document) -> None:
        """Agrega solo la palabra ATENTAMENTE centrada con espaciado."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("ATENTAMENTE.")
        run.bold = True
        run.font.size = Pt(11)

    def _add_firmante_nombre(self, doc: Document, nombre: str, cargo: str) -> None:
        """Agrega nombre y cargo del firmante centrados."""
        # Nombre
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(0)
        run_name = p_name.add_run(f"{nombre.upper()}.")
        run_name.bold = True
        run_name.font.size = Pt(11)

        # Cargo
        p_cargo = doc.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cargo.paragraph_format.space_after = Pt(0)
        run_cargo = p_cargo.add_run(f"{cargo.upper()}.")
        run_cargo.bold = True
        run_cargo.font.size = Pt(11)

    def _add_copias(self, doc: Document, copias: list[str]) -> None:
        """Bloque de copias institucionales (c.c.p.)."""
        for i, copia in enumerate(copias):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            if i == 0:
                prefix = "c.c.p. "
            else:
                prefix = "       "  # Indentación para alinear con primera copia
            run = p.add_run(f"{prefix}{copia}")
            run.font.size = Pt(7)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_referencia(
        self, doc: Document,
        elaboro: Optional[str], reviso: Optional[str],
    ) -> None:
        """Referencia interna: iniciales Director / elaboró / revisó."""
        from app.services.correspondencia_service import generar_referencia_oficio
        self._add_empty_lines(doc, 1)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        ref = generar_referencia_oficio("DIR", elaboro, reviso)
        run = p.add_run(ref)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_sello_digital(
        self,
        doc: Document,
        qr_png_bytes: bytes,
        nombre_firmante: str,
        cargo_firmante: str,
        rfc_firmante: str,
        serial_certificado: str,
        fecha_firma: str,
        correo_firmante: str = "",
        folio_firma: str = "",
    ) -> None:
        """
        Agrega bloque de firma electrónica institucional.
        Diseño: línea separadora + tabla [QR | Datos] con borde sutil.
        El nombre y cargo ya están arriba de este bloque.
        """
        from datetime import datetime

        # --- Línea separadora delgada gris ---
        self._add_empty_lines(doc, 1)
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(4)
        p_sep.paragraph_format.space_after = Pt(6)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(
            qn("w:bottom"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "CCCCCC"},
        )
        pBdr.append(bottom)
        pPr.append(pBdr)

        # --- Tabla 2 columnas: QR | Datos de firma ---
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Eliminar bordes predeterminados de la tabla
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_el = tblBorders.makeelement(
                qn(f"w:{border_name}"),
                {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0",
                 qn("w:color"): "DDDDDD" if border_name in ("insideV",) else "CCCCCC"},
            )
            tblBorders.append(border_el)
        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        # Ancho de columnas: QR ~4cm | Datos ~12cm
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(12)
        row = table.rows[0]

        # --- Columna izquierda: imagen QR (más grande ~3.5cm) ---
        cell_qr = row.cells[0]
        cell_qr.width = Cm(4)
        p_qr = cell_qr.paragraphs[0]
        p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Centrar verticalmente la celda del QR
        tc_pr = cell_qr._tc.get_or_add_tcPr()
        v_align = tc_pr.makeelement(qn("w:vAlign"), {qn("w:val"): "center"})
        tc_pr.append(v_align)
        if qr_png_bytes:
            run_qr = p_qr.add_run()
            qr_stream = io.BytesIO(qr_png_bytes)
            run_qr.add_picture(qr_stream, width=Cm(3.5))

        # --- Columna derecha: datos de firma ---
        cell_meta = row.cells[1]
        cell_meta.width = Cm(12)
        # Centrar verticalmente
        tc_pr2 = cell_meta._tc.get_or_add_tcPr()
        v_align2 = tc_pr2.makeelement(qn("w:vAlign"), {qn("w:val"): "center"})
        tc_pr2.append(v_align2)

        # Formatear fecha
        try:
            dt = datetime.fromisoformat(fecha_firma.replace("Z", "+00:00"))
            fecha_fmt = dt.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            fecha_fmt = fecha_firma

        # Datos de firma como pares etiqueta/valor
        firma_fields = [
            ("Firmado electrónicamente por:", nombre_firmante.upper()),
            ("No.Certificado:", serial_certificado or "—"),
            ("Fecha:", fecha_fmt),
            ("Folio:", folio_firma or "—"),
            ("Correo electrónico:", correo_firmante or "—"),
        ]

        first = True
        for label, value in firma_fields:
            if first:
                p_meta = cell_meta.paragraphs[0]
                first = False
            else:
                p_meta = cell_meta.add_paragraph()

            p_meta.paragraph_format.space_after = Pt(1)
            p_meta.paragraph_format.space_before = Pt(3)

            # Etiqueta en bold
            run_label = p_meta.add_run(label + "\n")
            run_label.bold = True
            run_label.font.size = Pt(7)
            run_label.font.name = "Arial"
            run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # Valor
            run_value = p_meta.add_run(value)
            run_value.bold = False
            run_value.font.size = Pt(7)
            run_value.font.name = "Arial"
            run_value.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


oficio_generator = OficioGeneratorService()
